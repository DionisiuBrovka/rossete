"""
DOCX exporter for Rossete.
"""

import logging
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..blocks import BlockContent
from ..exceptions import ExportError
from .base import Exporter

logger = logging.getLogger(__name__)


class DocxExporter(Exporter):
    """Exports document to DOCX format."""

    @property
    def format_name(self) -> str:
        """Return format name."""
        return "docx"

    def export(self, blocks: List[BlockContent], output_path: str) -> bool:
        """
        Export blocks to DOCX file.

        Args:
            blocks: List of BlockContent objects
            output_path: Path for output file

        Returns:
            True if successful

        Raises:
            ExportError: If export fails
        """
        try:
            if not self.validate_blocks(blocks):
                raise ExportError("No valid blocks to export")

            logger.debug(f"Exporting {len(blocks)} blocks to DOCX: {output_path}")

            # Create document
            doc = Document()
            self._setup_document_style(doc)

            # Add blocks
            for block in blocks:
                self._add_block_to_document(doc, block)

            # Save document
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_file)

            logger.info(f"Successfully exported to DOCX: {output_path}")
            return True

        except Exception as e:
            logger.error(f"DOCX export error: {e}")
            raise ExportError(f"DOCX export failed: {e}") from e

    def _setup_document_style(self, doc: Document) -> None:
        """Setup document formatting."""
        # Set default font for all styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        # Setup margins
        sections = doc.sections
        for section in sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = (
                Pt(72 * 3 / 4)  # 0.75 inches / ~19mm
            )

    def _add_block_to_document(self, doc: Document, block: BlockContent) -> None:
        """
        Add block content to document.

        Args:
            doc: Document object
            block: BlockContent to add
        """
        # Parse text as paragraphs (simple approach)
        lines = block.text.split("\n")

        for line in lines:
            if not line.strip():
                # Add blank paragraph
                doc.add_paragraph("")
                continue

            # Check if it's a heading (starts with #)
            if line.startswith("##"):
                level = 2
                text = line.lstrip("#").strip()
                paragraph = doc.add_heading(text, level=level)
            elif line.startswith("#"):
                level = 1
                text = line.lstrip("#").strip()
                paragraph = doc.add_heading(text, level=level)
            elif line.startswith("- "):
                # List item
                text = line[2:].strip()
                paragraph = doc.add_paragraph(text, style="List Bullet")
            else:
                # Regular paragraph
                paragraph = doc.add_paragraph(line)

            # Apply formatting if specified
            if block.formatting:
                self._apply_formatting(paragraph, block.formatting)

    def _apply_formatting(self, paragraph, formatting: dict) -> None:
        """
        Apply formatting to paragraph.

        Args:
            paragraph: Paragraph object
            formatting: Formatting dictionary
        """
        if not formatting:
            return

        # Handle bold
        if formatting.get("bold"):
            for run in paragraph.runs:
                run.bold = True

        # Handle italic
        if formatting.get("italic"):
            for run in paragraph.runs:
                run.italic = True

        # Handle alignment
        if "alignment" in formatting:
            alignment = formatting["alignment"]
            if alignment == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
